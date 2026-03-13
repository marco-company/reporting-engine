# Copyright 2026 Hunki Enterprises BV
# License AGPL-3.0 or later (https://www.gnu.org/licenses/agpl-3.0)

from io import BytesIO
from zipfile import ZipFile

from docx import Document

from odoo.tests.common import Form, TransactionCase


class TestReportDocx(TransactionCase):
    def test_demo_reports(self):
        zip_buffer, ext = self.env["ir.actions.report"]._render(
            "ir_module_multi_mode_zip",
            (
                self.env.ref("base.module_report_docx")
                + self.env.ref("base.module_base")
            ).ids,
        )
        self.assertEqual(ext, "zip")
        with ZipFile(BytesIO(zip_buffer)) as zip_file:
            self.assertItemsEqual(
                zip_file.namelist(), ["report_docx.docx", "base.docx"]
            )
        docx_buffer, ext = self.env["ir.actions.report"]._render(
            "ir_module_multi_mode_template", self.env.ref("base.module_base").ids
        )
        self.assertEqual(ext, "docx")
        all_text = "\n".join(p.text for p in Document(BytesIO(docx_buffer)).paragraphs)
        self.assertIn("That’s not so many modules", all_text)
        docx_buffer, ext = self.env["ir.actions.report"]._render(
            "ir_module_multi_mode_template",
            self.env["ir.module.module"].search([], limit=40).ids,
        )
        all_text = "\n".join(p.text for p in Document(BytesIO(docx_buffer)).paragraphs)
        self.assertIn("This are a lot of modules!", all_text)
        self.assertNotIn("This are a looooot of modules!", all_text)

    def test_form(self):
        with Form(
            self.env.ref("report_docx.report_ir_module_multi_mode_template")
        ) as report_form:
            report_form.docx_expression_test_record = self.env.ref("base.module_base")
            self.assertIn("format_amount", report_form.docx_help)
            report_form.docx_expression_test_expression = "object.nonexisting"
            self.assertIn(
                "has no attribute 'nonexisting'",
                report_form.docx_expression_test_result,
            )
            report_form.docx_expression_test_expression = "object.name"
            self.assertEqual(report_form.docx_expression_test_result, "base")
            self.assertEqual(report_form.docx_expression_test_code, "{{ object.name }}")
